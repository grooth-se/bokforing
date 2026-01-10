"""
Rapportgenerering - Genererar dokument från mallar

Stödjer generering av:
- Årsredovisning
- Resultaträkning
- Balansräkning
- Aktiebok
- Styrelseprotokoll
- Bolagsstämmoprotokoll

Mallar lagras i: templates/
    /arsredovisning/
        k2_arsredovisning.html
        k3_arsredovisning.html
    /rapporter/
        resultatrakning.html
        balansrakning.html
    /protokoll/
        styrelsemote.html
        bolagsstamma.html
    /register/
        aktiebok.html
"""
import os
from datetime import date, datetime
from decimal import Decimal
from typing import Dict, List, Optional
from pathlib import Path
from jinja2 import Environment, FileSystemLoader, select_autoescape
from sqlalchemy.orm import Session
import base64

from app.models import Company, Account, FiscalYear
from app.services.accounting import AccountingService


class ReportGenerator:
    """
    Genererar dokument från Jinja2-mallar

    Mallstruktur:
    - Mallar är HTML-filer med Jinja2-syntax
    - Kan konverteras till PDF med weasyprint eller liknande
    - Innehåller platshållare för företagsdata
    """

    TEMPLATE_DIR = Path(__file__).parent.parent.parent / "templates"

    # Malltyper och deras platser
    TEMPLATE_TYPES = {
        'annual_report_k2': 'arsredovisning/k2_arsredovisning.html',
        'annual_report_k3': 'arsredovisning/k3_arsredovisning.html',
        'income_statement': 'rapporter/resultatrakning.html',
        'balance_sheet': 'rapporter/balansrakning.html',
        'trial_balance': 'rapporter/rabalans.html',
        'general_ledger': 'rapporter/huvudbok.html',
        'shareholder_register': 'register/aktiebok.html',
        'board_meeting': 'protokoll/styrelsemote.html',
        'annual_meeting': 'protokoll/bolagsstamma.html',
    }

    def __init__(self, db: Session):
        self.db = db
        self.accounting_service = AccountingService(db)

        # Skapa mallmapp om den inte finns
        self.TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

        # Initiera Jinja2
        self.env = Environment(
            loader=FileSystemLoader(str(self.TEMPLATE_DIR)),
            autoescape=select_autoescape(['html', 'xml'])
        )

        # Lägg till filter
        self.env.filters['currency'] = self._currency_filter
        self.env.filters['date_format'] = self._date_filter

    def _currency_filter(self, value) -> str:
        """Formatera tal som valuta"""
        if value is None:
            return "0 kr"
        try:
            num = float(value)
            return f"{num:,.0f} kr".replace(",", " ")
        except (ValueError, TypeError):
            return str(value)

    def _date_filter(self, value, format_str: str = "%Y-%m-%d") -> str:
        """Formatera datum"""
        if value is None:
            return ""
        if isinstance(value, str):
            return value
        return value.strftime(format_str)

    def get_available_templates(self) -> Dict[str, bool]:
        """Lista tillgängliga mallar"""
        available = {}
        for name, path in self.TEMPLATE_TYPES.items():
            template_path = self.TEMPLATE_DIR / path
            available[name] = template_path.exists()
        return available

    def generate_annual_report(
        self,
        company_id: int,
        fiscal_year_id: int,
        additional_data: Dict = None
    ) -> str:
        """
        Generera årsredovisning

        Args:
            company_id: Företags-ID
            fiscal_year_id: Räkenskapsår-ID
            additional_data: Extra data (styrelseledamöter, etc.)

        Returns:
            HTML-sträng
        """
        company = self.db.query(Company).filter(Company.id == company_id).first()
        fiscal_year = self.db.query(FiscalYear).filter(FiscalYear.id == fiscal_year_id).first()

        if not company or not fiscal_year:
            raise ValueError("Företag eller räkenskapsår finns inte")

        # Välj mall baserat på K2/K3
        template_key = f"annual_report_{company.accounting_standard.value.lower()}"
        template_path = self.TEMPLATE_TYPES.get(template_key)

        if not template_path or not (self.TEMPLATE_DIR / template_path).exists():
            # Använd standardmall
            return self._generate_default_annual_report(company, fiscal_year, additional_data)

        # Hämta finansiell data
        financial_data = self._get_financial_data(company_id, fiscal_year)

        # Bygg kontext
        context = {
            'company': company,
            'fiscal_year': fiscal_year,
            'generated_at': datetime.now(),
            **financial_data,
            **(additional_data or {})
        }

        # Lägg till logotyp om den finns
        if company.logo:
            logo_b64 = base64.b64encode(company.logo).decode('utf-8')
            context['logo_base64'] = f"data:{company.logo_mimetype};base64,{logo_b64}"

        template = self.env.get_template(template_path)
        return template.render(context)

    def _get_financial_data(self, company_id: int, fiscal_year: FiscalYear) -> Dict:
        """Hämta finansiell data för rapporter"""
        # Råbalans
        trial_balance = self.accounting_service.get_trial_balance(
            company_id, fiscal_year.end_date
        )

        # Gruppera per kontoklass - resultaträkning
        income_statement = {
            'revenue': [],
            'expenses': [],
            'goods_cost': [],
            'other_costs': [],
            'personnel': [],
            'financial': [],
            'total_revenue': Decimal(0),
            'total_expenses': Decimal(0),
            'total_goods_cost': Decimal(0),
            'total_other_costs': Decimal(0),
            'total_personnel': Decimal(0),
            'total_financial': Decimal(0),
        }

        # Gruppera per kontoklass - balansräkning (detaljerad)
        balance_sheet = {
            'assets': [],
            'liabilities': [],
            'fixed_assets': [],          # 10xx-13xx Anläggningstillgångar
            'current_assets': [],        # 14xx-19xx Omsättningstillgångar
            'equity': [],                # 20xx-21xx Eget kapital
            'long_term_liabilities': [], # 22xx-24xx Långfristiga skulder
            'short_term_liabilities': [], # 25xx-29xx Kortfristiga skulder
            'total_assets': Decimal(0),
            'total_liabilities': Decimal(0),
            'total_fixed_assets': Decimal(0),
            'total_current_assets': Decimal(0),
            'total_equity': Decimal(0),
            'total_long_term': Decimal(0),
            'total_short_term': Decimal(0),
        }

        for account_data in trial_balance:
            number = account_data['account_number']
            balance = Decimal(str(account_data['balance']))

            # Tillgångar (1xxx)
            if number.startswith('1'):
                balance_sheet['assets'].append(account_data)
                balance_sheet['total_assets'] += balance

                # Anläggningstillgångar (10xx-13xx)
                if number.startswith(('10', '11', '12', '13')):
                    balance_sheet['fixed_assets'].append(account_data)
                    balance_sheet['total_fixed_assets'] += balance
                # Omsättningstillgångar (14xx-19xx)
                else:
                    balance_sheet['current_assets'].append(account_data)
                    balance_sheet['total_current_assets'] += balance

            # Eget kapital och skulder (2xxx)
            elif number.startswith('2'):
                balance_sheet['liabilities'].append(account_data)
                balance_sheet['total_liabilities'] += balance

                # Eget kapital (20xx-21xx)
                if number.startswith(('20', '21')):
                    balance_sheet['equity'].append(account_data)
                    balance_sheet['total_equity'] += balance
                # Långfristiga skulder (22xx-24xx)
                elif number.startswith(('22', '23', '24')):
                    balance_sheet['long_term_liabilities'].append(account_data)
                    balance_sheet['total_long_term'] += balance
                # Kortfristiga skulder (25xx-29xx)
                else:
                    balance_sheet['short_term_liabilities'].append(account_data)
                    balance_sheet['total_short_term'] += balance

            # Intäkter (3xxx)
            elif number.startswith('3'):
                income_statement['revenue'].append(account_data)
                income_statement['total_revenue'] += abs(balance)

            # Kostnader (4xxx-8xxx)
            elif number[0] in '45678':
                income_statement['expenses'].append(account_data)
                income_statement['total_expenses'] += balance

                # Varukostnad (4xxx)
                if number.startswith('4'):
                    income_statement['goods_cost'].append(account_data)
                    income_statement['total_goods_cost'] += balance
                # Övriga kostnader (5xxx-6xxx)
                elif number.startswith(('5', '6')):
                    income_statement['other_costs'].append(account_data)
                    income_statement['total_other_costs'] += balance
                # Personalkostnader (7xxx)
                elif number.startswith('7'):
                    income_statement['personnel'].append(account_data)
                    income_statement['total_personnel'] += balance
                # Finansiella poster (8xxx)
                elif number.startswith('8'):
                    income_statement['financial'].append(account_data)
                    income_statement['total_financial'] += balance

        # Resultat
        result = income_statement['total_revenue'] - income_statement['total_expenses']

        # Bruttovinst och rörelseresultat
        income_statement['gross_profit'] = income_statement['total_revenue'] - income_statement['total_goods_cost']
        income_statement['operating_result'] = (
            income_statement['gross_profit'] -
            income_statement['total_other_costs'] -
            income_statement['total_personnel']
        )
        income_statement['result'] = result

        return {
            'trial_balance': trial_balance,
            'income_statement': income_statement,
            'balance_sheet': balance_sheet,
            'result': result,
        }

    def _generate_default_annual_report(
        self,
        company: Company,
        fiscal_year: FiscalYear,
        additional_data: Dict = None
    ) -> str:
        """Generera enkel årsredovisning utan mall"""
        financial_data = self._get_financial_data(company.id, fiscal_year)

        # Bygg HTML
        html = f"""
        <!DOCTYPE html>
        <html lang="sv">
        <head>
            <meta charset="UTF-8">
            <title>Årsredovisning {company.name} {fiscal_year.end_date.year}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                h1 {{ color: #333; border-bottom: 2px solid #333; padding-bottom: 10px; }}
                h2 {{ color: #555; margin-top: 30px; }}
                table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                th, td {{ padding: 8px 12px; text-align: left; border-bottom: 1px solid #ddd; }}
                th {{ background-color: #f5f5f5; }}
                .amount {{ text-align: right; }}
                .total {{ font-weight: bold; border-top: 2px solid #333; }}
                .header {{ text-align: center; margin-bottom: 40px; }}
                .logo {{ max-width: 150px; margin-bottom: 20px; }}
                .company-info {{ margin-bottom: 30px; }}
                .footer {{ margin-top: 50px; font-size: 0.9em; color: #666; }}
            </style>
        </head>
        <body>
            <div class="header">
                {"<img src='data:" + company.logo_mimetype + ";base64," + base64.b64encode(company.logo).decode() + "' class='logo'>" if company.logo else ""}
                <h1>Årsredovisning</h1>
                <h2>{company.name}</h2>
                <p>Organisationsnummer: {company.org_number}</p>
                <p>Räkenskapsår: {fiscal_year.start_date} - {fiscal_year.end_date}</p>
            </div>

            <h2>Resultaträkning</h2>
            <table>
                <tr><th>Konto</th><th>Namn</th><th class="amount">Belopp</th></tr>
                <tr><td colspan="3"><strong>Intäkter</strong></td></tr>
        """

        for item in financial_data['income_statement']['revenue']:
            html += f"<tr><td>{item['account_number']}</td><td>{item['account_name']}</td><td class='amount'>{abs(item['balance']):,.0f} kr</td></tr>"

        html += f"""
                <tr class="total"><td></td><td>Summa intäkter</td><td class="amount">{financial_data['income_statement']['total_revenue']:,.0f} kr</td></tr>
                <tr><td colspan="3"><strong>Kostnader</strong></td></tr>
        """

        for item in financial_data['income_statement']['expenses']:
            html += f"<tr><td>{item['account_number']}</td><td>{item['account_name']}</td><td class='amount'>{item['balance']:,.0f} kr</td></tr>"

        html += f"""
                <tr class="total"><td></td><td>Summa kostnader</td><td class="amount">{financial_data['income_statement']['total_expenses']:,.0f} kr</td></tr>
                <tr class="total"><td></td><td><strong>Årets resultat</strong></td><td class="amount"><strong>{financial_data['result']:,.0f} kr</strong></td></tr>
            </table>

            <h2>Balansräkning</h2>
            <table>
                <tr><th>Konto</th><th>Namn</th><th class="amount">Belopp</th></tr>
                <tr><td colspan="3"><strong>Tillgångar</strong></td></tr>
        """

        for item in financial_data['balance_sheet']['assets']:
            if item['balance'] != 0:
                html += f"<tr><td>{item['account_number']}</td><td>{item['account_name']}</td><td class='amount'>{item['balance']:,.0f} kr</td></tr>"

        html += f"""
                <tr class="total"><td></td><td>Summa tillgångar</td><td class="amount">{financial_data['balance_sheet']['total_assets']:,.0f} kr</td></tr>
                <tr><td colspan="3"><strong>Eget kapital och skulder</strong></td></tr>
        """

        for item in financial_data['balance_sheet']['liabilities']:
            if item['balance'] != 0:
                html += f"<tr><td>{item['account_number']}</td><td>{item['account_name']}</td><td class='amount'>{abs(item['balance']):,.0f} kr</td></tr>"

        html += f"""
                <tr class="total"><td></td><td>Summa eget kapital och skulder</td><td class="amount">{abs(financial_data['balance_sheet']['total_liabilities']):,.0f} kr</td></tr>
            </table>

            <div class="footer">
                <p>Genererad: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
                <p>Redovisningsstandard: {company.accounting_standard.value}</p>
            </div>
        </body>
        </html>
        """

        return html

    def generate_income_statement(
        self,
        company_id: int,
        fiscal_year_id: int
    ) -> str:
        """Generera resultaträkning"""
        company = self.db.query(Company).filter(Company.id == company_id).first()
        fiscal_year = self.db.query(FiscalYear).filter(FiscalYear.id == fiscal_year_id).first()
        financial_data = self._get_financial_data(company_id, fiscal_year)

        # Kontrollera om mall finns
        template_path = self.TEMPLATE_TYPES.get('income_statement')
        if template_path and (self.TEMPLATE_DIR / template_path).exists():
            template = self.env.get_template(template_path)
            return template.render(
                company=company,
                fiscal_year=fiscal_year,
                generated_at=datetime.now(),
                show_previous=False,
                **financial_data
            )

        # Standardrapport
        return self._generate_simple_report(
            "Resultaträkning",
            company,
            fiscal_year,
            financial_data['income_statement']
        )

    def generate_balance_sheet(
        self,
        company_id: int,
        fiscal_year_id: int
    ) -> str:
        """Generera balansräkning"""
        company = self.db.query(Company).filter(Company.id == company_id).first()
        fiscal_year = self.db.query(FiscalYear).filter(FiscalYear.id == fiscal_year_id).first()
        financial_data = self._get_financial_data(company_id, fiscal_year)

        # Kontrollera om mall finns
        template_path = self.TEMPLATE_TYPES.get('balance_sheet')
        if template_path and (self.TEMPLATE_DIR / template_path).exists():
            template = self.env.get_template(template_path)
            return template.render(
                company=company,
                fiscal_year=fiscal_year,
                generated_at=datetime.now(),
                show_previous=False,
                **financial_data
            )

        return self._generate_simple_report(
            "Balansräkning",
            company,
            fiscal_year,
            financial_data['balance_sheet']
        )

    def generate_shareholder_register(
        self,
        company_id: int,
        shareholders: List[Dict]
    ) -> str:
        """Generera aktiebok"""
        company = self.db.query(Company).filter(Company.id == company_id).first()

        template_path = self.TEMPLATE_TYPES.get('shareholder_register')
        if template_path and (self.TEMPLATE_DIR / template_path).exists():
            template = self.env.get_template(template_path)
            return template.render(
                company=company,
                shareholders=shareholders,
                generated_at=datetime.now()
            )

        # Standardaktiebook
        return self._generate_default_shareholder_register(company, shareholders)

    def _generate_default_shareholder_register(
        self,
        company: Company,
        shareholders: List[Dict]
    ) -> str:
        """Generera enkel aktiebok"""
        html = f"""
        <!DOCTYPE html>
        <html lang="sv">
        <head>
            <meta charset="UTF-8">
            <title>Aktiebok - {company.name}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ border-bottom: 2px solid #333; padding-bottom: 10px; }}
                table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                th, td {{ padding: 10px; border: 1px solid #ddd; }}
                th {{ background-color: #f5f5f5; }}
            </style>
        </head>
        <body>
            <h1>Aktiebok</h1>
            <p><strong>{company.name}</strong></p>
            <p>Org.nr: {company.org_number}</p>

            <table>
                <tr>
                    <th>Aktienummer</th>
                    <th>Ägare</th>
                    <th>Personnr/Orgnr</th>
                    <th>Antal aktier</th>
                    <th>Förvärvsdag</th>
                </tr>
        """

        total_shares = 0
        for sh in shareholders:
            html += f"""
                <tr>
                    <td>{sh.get('share_numbers', '-')}</td>
                    <td>{sh.get('name', '')}</td>
                    <td>{sh.get('id_number', '')}</td>
                    <td>{sh.get('num_shares', 0)}</td>
                    <td>{sh.get('acquisition_date', '')}</td>
                </tr>
            """
            total_shares += sh.get('num_shares', 0)

        html += f"""
                <tr>
                    <td colspan="3"><strong>Totalt</strong></td>
                    <td><strong>{total_shares}</strong></td>
                    <td></td>
                </tr>
            </table>
            <p>Genererad: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
        </body>
        </html>
        """

        return html

    def _generate_simple_report(
        self,
        title: str,
        company: Company,
        fiscal_year: FiscalYear,
        data: Dict
    ) -> str:
        """Generera enkel rapport"""
        html = f"""
        <!DOCTYPE html>
        <html lang="sv">
        <head>
            <meta charset="UTF-8">
            <title>{title} - {company.name}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ border-bottom: 2px solid #333; }}
                table {{ width: 100%; border-collapse: collapse; }}
                th, td {{ padding: 8px; text-align: left; border-bottom: 1px solid #ddd; }}
                .amount {{ text-align: right; }}
            </style>
        </head>
        <body>
            <h1>{title}</h1>
            <p>{company.name} | {fiscal_year.start_date} - {fiscal_year.end_date}</p>
            <pre>{data}</pre>
        </body>
        </html>
        """
        return html

    def create_default_templates(self):
        """Skapa standardmallar i templates-mappen"""
        # Skapa mappstruktur
        for folder in ['arsredovisning', 'rapporter', 'protokoll', 'register']:
            (self.TEMPLATE_DIR / folder).mkdir(parents=True, exist_ok=True)

        # Skapa README
        readme_path = self.TEMPLATE_DIR / "README.md"
        if not readme_path.exists():
            readme_path.write_text("""# Dokumentmallar

Denna mapp innehåller Jinja2-mallar för generering av rapporter och dokument.

## Mappstruktur

- `arsredovisning/` - Mallar för årsredovisning (K2/K3)
- `rapporter/` - Finansiella rapporter
- `protokoll/` - Mötes- och styrelseprotokoll
- `register/` - Aktiebok och andra register

## Använda variabler

### Företagsdata
- `{{ company.name }}` - Företagsnamn
- `{{ company.org_number }}` - Organisationsnummer
- `{{ company.logo_base64 }}` - Logotyp som base64 (för <img>)

### Räkenskapsår
- `{{ fiscal_year.start_date }}` - Startdatum
- `{{ fiscal_year.end_date }}` - Slutdatum

### Finansiell data
- `{{ income_statement }}` - Resultaträkningsdata
- `{{ balance_sheet }}` - Balansräkningsdata
- `{{ trial_balance }}` - Råbalans

### Filter
- `{{ value|currency }}` - Formaterar som valuta
- `{{ date|date_format }}` - Formaterar datum

## Exempel

```html
<h1>Årsredovisning {{ company.name }}</h1>
<p>Org.nr: {{ company.org_number }}</p>
<p>Resultat: {{ result|currency }}</p>
```
""", encoding='utf-8')

    def to_pdf(self, html_content: str) -> bytes:
        """
        Konvertera HTML till PDF

        Försöker först WeasyPrint, sedan ReportLab som fallback.
        """
        # Försök WeasyPrint först
        try:
            from weasyprint import HTML
            pdf_bytes = HTML(string=html_content, base_url=str(self.TEMPLATE_DIR)).write_pdf()
            return pdf_bytes
        except (ImportError, OSError):
            pass  # Fallback till ReportLab

        # Fallback: Använd ReportLab för enkel PDF
        try:
            return self._generate_pdf_with_reportlab(html_content)
        except Exception as e:
            raise RuntimeError(f"Kunde inte generera PDF: {str(e)}")

    def _generate_pdf_with_reportlab(self, html_content: str) -> bytes:
        """Generera PDF med ReportLab från HTML-innehåll"""
        import io
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors

        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise RuntimeError("BeautifulSoup krävs. Kör: pip install beautifulsoup4")

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Swedish', fontName='Helvetica', fontSize=10, leading=14))
        styles.add(ParagraphStyle(name='SwedishTitle', fontName='Helvetica-Bold', fontSize=16, leading=20))
        styles.add(ParagraphStyle(name='SwedishHeading', fontName='Helvetica-Bold', fontSize=12, leading=16))

        story = []
        soup = BeautifulSoup(html_content, 'html.parser')

        # Extrahera rubriker och text
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'table']):
            if element.name == 'h1':
                story.append(Paragraph(element.get_text(strip=True), styles['SwedishTitle']))
                story.append(Spacer(1, 0.5*cm))
            elif element.name in ['h2', 'h3']:
                story.append(Spacer(1, 0.3*cm))
                story.append(Paragraph(element.get_text(strip=True), styles['SwedishHeading']))
                story.append(Spacer(1, 0.2*cm))
            elif element.name == 'p':
                text = element.get_text(strip=True)
                if text:
                    story.append(Paragraph(text, styles['Swedish']))
                    story.append(Spacer(1, 0.2*cm))
            elif element.name == 'table':
                # Konvertera tabell
                table_data = []
                for row in element.find_all('tr'):
                    row_data = []
                    for cell in row.find_all(['th', 'td']):
                        row_data.append(cell.get_text(strip=True))
                    if row_data:
                        table_data.append(row_data)

                if table_data:
                    # Skapa tabell med stil
                    t = Table(table_data)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 0.3*cm))

        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    def to_docx(self, html_content: str, title: str = "Rapport") -> bytes:
        """
        Konvertera HTML till DOCX

        Args:
            html_content: HTML-sträng
            title: Dokumenttitel

        Returns:
            DOCX som bytes
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
        except ImportError:
            raise RuntimeError("python-docx är inte installerat. Kör: pip install python-docx")

        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise RuntimeError("BeautifulSoup är inte installerat. Kör: pip install beautifulsoup4")

        # Skapa Word-dokument
        doc = Document()

        # Parsa HTML
        soup = BeautifulSoup(html_content, 'html.parser')

        # Extrahera titel från h1 om det finns
        h1 = soup.find('h1')
        if h1:
            heading = doc.add_heading(h1.get_text(strip=True), 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Gå igenom elementen
        for element in soup.find_all(['h2', 'h3', 'p', 'table']):
            if element.name == 'h2':
                doc.add_heading(element.get_text(strip=True), 1)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(strip=True), 2)
            elif element.name == 'p':
                text = element.get_text(strip=True)
                if text:
                    doc.add_paragraph(text)
            elif element.name == 'table':
                # Hantera tabeller
                rows = element.find_all('tr')
                if rows:
                    # Räkna max antal kolumner
                    max_cols = 0
                    for row in rows:
                        cols = len(row.find_all(['th', 'td']))
                        max_cols = max(max_cols, cols)

                    if max_cols > 0:
                        table = doc.add_table(rows=0, cols=max_cols)
                        table.style = 'Table Grid'

                        for row in rows:
                            cells = row.find_all(['th', 'td'])
                            doc_row = table.add_row()
                            for i, cell in enumerate(cells):
                                if i < max_cols:
                                    doc_row.cells[i].text = cell.get_text(strip=True)

                        doc.add_paragraph()  # Lägg till mellanrum efter tabell

        # Spara till bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def generate_report_with_export(
        self,
        report_type: str,
        company_id: int,
        fiscal_year_id: int,
        output_format: str = "html",
        **kwargs
    ) -> tuple:
        """
        Generera rapport med val av exportformat

        Args:
            report_type: Rapporttyp ('annual_report', 'income_statement', 'balance_sheet', etc.)
            company_id: Företags-ID
            fiscal_year_id: Räkenskapsår-ID
            output_format: 'html', 'pdf', eller 'docx'
            **kwargs: Extra data

        Returns:
            tuple: (data_bytes, content_type, filename)
        """
        company = self.db.query(Company).filter(Company.id == company_id).first()
        fiscal_year = self.db.query(FiscalYear).filter(FiscalYear.id == fiscal_year_id).first()

        if not company or not fiscal_year:
            raise ValueError("Företag eller räkenskapsår finns inte")

        # Generera rapport baserat på typ
        if report_type == 'annual_report':
            html_content = self.generate_annual_report(company_id, fiscal_year_id, kwargs)
            base_filename = f"arsredovisning_{company.org_number}_{fiscal_year.end_date.year}"
        elif report_type == 'income_statement':
            html_content = self.generate_income_statement(company_id, fiscal_year_id)
            base_filename = f"resultatrakning_{company.org_number}_{fiscal_year.end_date.year}"
        elif report_type == 'balance_sheet':
            html_content = self.generate_balance_sheet(company_id, fiscal_year_id)
            base_filename = f"balansrakning_{company.org_number}_{fiscal_year.end_date.year}"
        elif report_type == 'trial_balance':
            html_content = self._generate_trial_balance_report(company_id, fiscal_year_id)
            base_filename = f"rabalans_{company.org_number}_{fiscal_year.end_date.year}"
        elif report_type == 'general_ledger':
            html_content = self._generate_general_ledger_report(company_id, fiscal_year_id, kwargs.get('account_filter'))
            base_filename = f"huvudbok_{company.org_number}_{fiscal_year.end_date.year}"
        elif report_type == 'shareholder_register':
            html_content = self.generate_shareholder_register(company_id, kwargs.get('shareholders', []))
            base_filename = f"aktiebok_{company.org_number}"
        else:
            raise ValueError(f"Okänd rapporttyp: {report_type}")

        # Konvertera till önskat format
        if output_format == "html":
            return html_content.encode('utf-8'), "text/html", f"{base_filename}.html"
        elif output_format == "pdf":
            pdf_data = self.to_pdf(html_content)
            return pdf_data, "application/pdf", f"{base_filename}.pdf"
        elif output_format == "docx":
            docx_data = self.to_docx(html_content, report_type)
            return docx_data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{base_filename}.docx"
        else:
            raise ValueError(f"Okänt format: {output_format}")

    def _generate_trial_balance_report(self, company_id: int, fiscal_year_id: int) -> str:
        """Generera råbalansrapport"""
        company = self.db.query(Company).filter(Company.id == company_id).first()
        fiscal_year = self.db.query(FiscalYear).filter(FiscalYear.id == fiscal_year_id).first()
        trial_balance = self.accounting_service.get_trial_balance(company_id, fiscal_year.end_date)

        # Kontrollera om mall finns
        template_path = self.TEMPLATE_TYPES.get('trial_balance')
        if template_path and (self.TEMPLATE_DIR / template_path).exists():
            template = self.env.get_template(template_path)
            return template.render(
                company=company,
                fiscal_year=fiscal_year,
                trial_balance=trial_balance,
                generated_at=datetime.now()
            )

        # Standardrapport
        total_debit = sum(item.get('debit', 0) for item in trial_balance)
        total_credit = sum(item.get('credit', 0) for item in trial_balance)

        html = f"""
        <!DOCTYPE html>
        <html lang="sv">
        <head>
            <meta charset="UTF-8">
            <title>Råbalans - {company.name}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #2c5282; border-bottom: 2px solid #2c5282; }}
                table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                th, td {{ padding: 8px 12px; text-align: left; border-bottom: 1px solid #e2e8f0; }}
                th {{ background-color: #edf2f7; }}
                .amount {{ text-align: right; font-family: monospace; }}
                .total {{ font-weight: bold; background-color: #e2e8f0; }}
            </style>
        </head>
        <body>
            <h1>Råbalans</h1>
            <p><strong>{company.name}</strong> | Org.nr: {company.org_number}</p>
            <p>Per: {fiscal_year.end_date}</p>

            <table>
                <thead>
                    <tr>
                        <th>Konto</th>
                        <th>Namn</th>
                        <th class="amount">Debet</th>
                        <th class="amount">Kredit</th>
                    </tr>
                </thead>
                <tbody>
        """

        for item in trial_balance:
            debit = item.get('debit', 0)
            credit = item.get('credit', 0)
            html += f"""
                    <tr>
                        <td>{item['account_number']}</td>
                        <td>{item['account_name']}</td>
                        <td class="amount">{debit:,.0f} kr</td>
                        <td class="amount">{credit:,.0f} kr</td>
                    </tr>
            """

        html += f"""
                    <tr class="total">
                        <td></td>
                        <td><strong>Summa</strong></td>
                        <td class="amount"><strong>{total_debit:,.0f} kr</strong></td>
                        <td class="amount"><strong>{total_credit:,.0f} kr</strong></td>
                    </tr>
                </tbody>
            </table>
            <p>Genererad: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
        </body>
        </html>
        """
        return html

    def _generate_general_ledger_report(
        self,
        company_id: int,
        fiscal_year_id: int,
        account_filter: str = None
    ) -> str:
        """Generera huvudboksrapport"""
        company = self.db.query(Company).filter(Company.id == company_id).first()
        fiscal_year = self.db.query(FiscalYear).filter(FiscalYear.id == fiscal_year_id).first()

        accounts = self.accounting_service.get_accounts(company_id)
        transactions = self.accounting_service.get_transactions(company_id, fiscal_year_id)

        html = f"""
        <!DOCTYPE html>
        <html lang="sv">
        <head>
            <meta charset="UTF-8">
            <title>Huvudbok - {company.name}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; font-size: 10pt; }}
                h1 {{ color: #2c5282; border-bottom: 2px solid #2c5282; }}
                h2 {{ color: #2c5282; margin-top: 30px; page-break-before: auto; }}
                table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
                th, td {{ padding: 6px 10px; text-align: left; border-bottom: 1px solid #e2e8f0; }}
                th {{ background-color: #edf2f7; }}
                .amount {{ text-align: right; font-family: monospace; }}
                .total {{ font-weight: bold; background-color: #bee3f8; }}
                .ib {{ font-style: italic; color: #666; }}
            </style>
        </head>
        <body>
            <h1>Huvudbok</h1>
            <p><strong>{company.name}</strong> | Org.nr: {company.org_number}</p>
            <p>Räkenskapsår: {fiscal_year.start_date} - {fiscal_year.end_date}</p>
        """

        from app.config import AccountType

        for account in accounts:
            if account_filter and not account.number.startswith(account_filter):
                continue

            # Hitta transaktioner för detta konto
            account_txs = []
            running_balance = account.opening_balance or Decimal(0)

            for tx in transactions:
                for line in tx.lines:
                    if line.account_id == account.id:
                        if account.account_type in [AccountType.ASSET, AccountType.EXPENSE]:
                            running_balance += line.debit - line.credit
                        else:
                            running_balance += line.credit - line.debit

                        account_txs.append({
                            'date': tx.transaction_date,
                            'ver': tx.verification_number,
                            'desc': tx.description,
                            'debit': line.debit,
                            'credit': line.credit,
                            'balance': running_balance
                        })

            if account_txs or (account.opening_balance and account.opening_balance != 0):
                html += f"""
                <h2>{account.number} {account.name}</h2>
                <table>
                    <thead>
                        <tr>
                            <th>Datum</th>
                            <th>Ver</th>
                            <th>Beskrivning</th>
                            <th class="amount">Debet</th>
                            <th class="amount">Kredit</th>
                            <th class="amount">Saldo</th>
                        </tr>
                    </thead>
                    <tbody>
                        <tr class="ib">
                            <td></td>
                            <td></td>
                            <td>Ingående balans</td>
                            <td class="amount"></td>
                            <td class="amount"></td>
                            <td class="amount">{account.opening_balance or 0:,.0f} kr</td>
                        </tr>
                """

                for tx in account_txs:
                    html += f"""
                        <tr>
                            <td>{tx['date']}</td>
                            <td>{tx['ver']}</td>
                            <td>{tx['desc']}</td>
                            <td class="amount">{tx['debit']:,.0f} kr</td>
                            <td class="amount">{tx['credit']:,.0f} kr</td>
                            <td class="amount">{tx['balance']:,.0f} kr</td>
                        </tr>
                    """

                html += f"""
                        <tr class="total">
                            <td></td>
                            <td></td>
                            <td>Utgående balans</td>
                            <td class="amount"></td>
                            <td class="amount"></td>
                            <td class="amount">{running_balance:,.0f} kr</td>
                        </tr>
                    </tbody>
                </table>
                """

        html += f"""
            <p>Genererad: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
        </body>
        </html>
        """
        return html
